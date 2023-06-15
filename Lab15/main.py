from support_function import read_file
import docx


def ModificationOfTheStringLength(message: str, container: str):
    message_binary = ''.join([bin(ord(i))[2:].rjust(8, '0') for i in message])

    # Вставка в контейнер
    list_container = [i for i in container.split()]
    container_modify = ' '.join([list_container[i] + ' ' if message_binary[i] == '1' else list_container[i]
                                 for i in range(len(message_binary))] + list_container[len(message_binary):])
    print(f'----- Вставка в контейнер -----\n'
          f'Сообщение:  {message}\n'
          f'Сообщение(2): {message_binary}\n'
          f'Контейнер:\n{container}\n'
          f'Модифицированное сообщение:\n{container_modify}\n')

    # Запись в docx
    doc_modify = docx.Document()
    doc_modify.add_paragraph(container_modify)
    doc_modify.save('result.docx')

    # Извлечение из контейнера
    doc_modify = docx.Document('result.docx')
    for paragraph_ in doc_modify.paragraphs:
        for run in paragraph_.runs:
            extracted_container = run.text

    list_extracted_container = ''.join(['1' if i == '' else '0'
                                        for i in extracted_container.split(' ')]).replace('01', '1')
    extracted_message = ''.join([chr(int(list_extracted_container[i * 8: (i + 1) * 8], 2))
                                 for i in range(len(list_extracted_container) // 8)]).rstrip('\x00')
    print(f'----- Извлечение из контейнера -----\n'
          f'Извлечённый контейнер: {extracted_container}\n'
          f'Извлечённое сообщение: {extracted_message}\n')


def main():
    container = read_file('text.txt', -1)
    message = 'Dashchinskii Maksim Leonidovich'

    print(f'----- Изменение длины строки -----')
    ModificationOfTheStringLength(message, container)


if __name__ == '__main__':
    main()
